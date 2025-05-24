from typing import Dict, List, TypedDict
from langgraph.graph import StateGraph, END
from langchain.prompts import PromptTemplate
from langchain_community.llms import HuggingFacePipeline
from transformers import pipeline
import pdfplumber
import json
import logging
import os
import re
import sys
import pytesseract
from PIL import Image
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

# Set up logging for debugging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

# Define the state to track data across the workflow
class RFPState(TypedDict):
    rfp_path: str
    rfp_text: str
    qa_pairs: List[Dict[str, str]]
    response: str

# Verify dependencies
try:
    import spacy
    nlp = spacy.load("en_core_web_sm")
    logger.info("spaCy model loaded successfully")
    use_spacy = True
except Exception as e:
    logger.warning(f"spaCy not available, using keyword-based extraction: {e}")
    use_spacy = False

try:
    llm = HuggingFacePipeline.from_model_id(
        model_id="google/flan-t5-small",
        task="text2text-generation",
        pipeline_kwargs={"max_length": 512, "truncation": True}
    )
    logger.info("LLM model loaded successfully")
except Exception as e:
    logger.error(f"Failed to load LLM: {e}")
    raise

# Node 1: Extract text from RFP PDF
def extract_rfp_text(state: RFPState) -> RFPState:
    logger.info(f"Attempting to extract text from RFP: {state['rfp_path']}")
    if not os.path.exists(state["rfp_path"]):
        error_msg = (
            f"RFP file not found at: {state['rfp_path']}\n"
            "Please ensure 'apexneural_rfp.pdf' exists in the project directory."
        )
        logger.error(error_msg)
        raise FileNotFoundError(error_msg)
    
    try:
        text = ""
        with pdfplumber.open(state["rfp_path"]) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                else:
                    # Fallback: Try OCR if text extraction fails (e.g., for scanned PDFs)
                    logger.info("No text extracted from page, attempting OCR...")
                    try:
                        # Convert page to image for OCR
                        page_image = page.to_image(resolution=300)
                        img_byte_arr = io.BytesIO()
                        page_image.original.save(img_byte_arr, format='PNG')
                        img_byte_arr.seek(0)
                        img = Image.open(img_byte_arr)
                        ocr_text = pytesseract.image_to_string(img)
                        if ocr_text:
                            text += ocr_text + "\n"
                            logger.info(f"OCR extracted text: {ocr_text[:100]}...")
                    except Exception as ocr_error:
                        logger.warning(f"OCR failed for page: {ocr_error}")
                        continue
        
        if not text.strip():
            logger.warning("No text extracted from RFP after attempting OCR.")
            state["rfp_text"] = ""
        else:
            logger.info(f"Extracted text: {text[:200]}...")
            state["rfp_text"] = text
            logger.info("RFP text extracted successfully")
    except Exception as e:
        logger.error(f"Error extracting RFP text: {e}")
        state["rfp_text"] = ""
    return state

# Node 2: Generate Q&A from RFP text
def generate_qa_pairs(state: RFPState) -> RFPState:
    logger.info("Generating Q&A pairs from RFP text")
    if not state["rfp_text"]:
        logger.warning("No RFP text available for Q&A generation")
        state["qa_pairs"] = []
        return state

    all_qa_pairs = []
    
    # Step 1: Extract potential requirements
    if use_spacy:
        doc = nlp(state["rfp_text"])
        requirements = [
            sent.text.strip() for sent in doc.sents
            if any(keyword in sent.text.lower() for keyword in ["must", "shall", "required", "scope", "resource", "location", "capacity", "timeline", "duration"])
        ]
    else:
        sentences = re.split(r'[.!?]\s+', state["rfp_text"])
        requirements = [
            sent.strip() for sent in sentences
            if any(keyword in sent.lower() for keyword in ["must", "shall", "required", "scope", "resource", "location", "capacity", "timeline", "duration"])
        ]
    
    # Step 2: Look for numbered clauses (e.g., "1.1.1") to capture more details
    clause_pattern = r'\d+\.\d+(\.\d+)?\s+.*?(?=\n\d+\.\d+|\n[A-Z]|\Z)'
    clauses = re.findall(clause_pattern, state["rfp_text"], re.DOTALL)
    requirements.extend([clause.strip() for clause in clauses if clause.strip()])
    
    # Remove duplicates while preserving order
    seen = set()
    requirements = [req for req in requirements if not (req in seen or seen.add(req))]
    
    logger.info(f"Found {len(requirements)} potential requirements: {requirements[:3]}...")

    # Step 3: Define multiple prompt templates for different types of questions
    qa_prompts = {
        "scope": PromptTemplate(
            input_variables=["requirement"],
            template="Generate a question and answer about the scope of work:\nRequirement: {requirement}\nQuestion: What is the scope of work related to this requirement?\nAnswer: {requirement}"
        ),
        "resource": PromptTemplate(
            input_variables=["requirement"],
            template="Generate a question and answer about the resources needed:\nRequirement: {requirement}\nQuestion: What resources are required for this?\nAnswer: {requirement}"
        ),
        "location": PromptTemplate(
            input_variables=["requirement"],
            template="Generate a question and answer about the work location:\nRequirement: {requirement}\nQuestion: Where must this work be performed?\nAnswer: {requirement}"
        ),
        "timeline": PromptTemplate(
            input_variables=["requirement"],
            template="Generate a question and answer about the timeline:\nRequirement: {requirement}\nQuestion: What is the timeline for this requirement?\nAnswer: {requirement}"
        ),
        "general": PromptTemplate(
            input_variables=["requirement"],
            template="Generate a question and answer:\nRequirement: {requirement}\nQuestion: What is the requirement for this?\nAnswer: {requirement}"
        )
    }
    
    # Step 4: Generate Q&A pairs for all requirements
    for idx, req in enumerate(requirements):
        try:
            # Clean up the requirement text
            cleaned_req = req.replace('\n-', '').strip()
            if not cleaned_req:
                continue
                
            # Determine the type of prompt to use based on keywords
            prompt_type = "general"
            if any(keyword in cleaned_req.lower() for keyword in ["scope", "work", "responsibility"]):
                prompt_type = "scope"
            elif any(keyword in cleaned_req.lower() for keyword in ["resource", "cost", "security", "deposit"]):
                prompt_type = "resource"
            elif any(keyword in cleaned_req.lower() for keyword in ["location", "place", "address", "grid"]):
                prompt_type = "location"
            elif any(keyword in cleaned_req.lower() for keyword in ["timeline", "duration", "date", "period"]):
                prompt_type = "timeline"

            prompt = qa_prompts[prompt_type].format(requirement=cleaned_req)
            qa_text = llm.invoke(prompt)
            logger.info(f"LLM output for requirement {idx + 1} ('{cleaned_req[:50]}...'): {qa_text}")
            
            # Parse the LLM output
            try:
                # Handle single-line or malformed outputs
                if "\n" in qa_text:
                    lines = qa_text.split("\n", 1)
                    question = lines[0].replace("Question: ", "").strip()
                    answer = lines[1].replace("Answer: ", "").strip() if len(lines) > 1 else cleaned_req
                else:
                    # If no newline, assume it's a question or use the prompt's question
                    question = qa_prompts[prompt_type].template.split("Question: ")[1].split("\n")[0].format(requirement=cleaned_req[:20])
                    answer = cleaned_req
                if question and answer:
                    all_qa_pairs.append({
                        "question": question,
                        "answer": answer,
                        "prompt_type": prompt_type,
                        "requirement": cleaned_req
                    })
                    logger.info(f"Generated Q&A {idx + 1}: Q: {question[:50]}... A: {answer[:50]}...")
                else:
                    logger.warning(f"Empty question or answer for requirement {idx + 1}: {cleaned_req[:50]}...")
                    # Fallback
                    question = f"What is the detail for {cleaned_req[:20].lower()}...?"
                    answer = cleaned_req
                    all_qa_pairs.append({
                        "question": question,
                        "answer": answer,
                        "prompt_type": prompt_type,
                        "requirement": cleaned_req
                    })
                    logger.info(f"Fallback Q&A {idx + 1}: Q: {question[:50]}... A: {answer[:50]}...")
            except Exception as parse_error:
                logger.warning(f"Failed to parse Q&A for requirement {idx + 1}: {cleaned_req[:50]}... {parse_error}")
                # Fallback
                question = f"What is the detail for {cleaned_req[:20].lower()}...?"
                answer = cleaned_req
                all_qa_pairs.append({
                    "question": question,
                    "answer": answer,
                    "prompt_type": prompt_type,
                    "requirement": cleaned_req
                })
                logger.info(f"Fallback Q&A {idx + 1}: Q: {question[:50]}... A: {answer[:50]}...")
        except Exception as e:
            logger.error(f"Error generating Q&A for requirement {idx + 1}: {cleaned_req[:50]}... {e}")
            continue

    # Step 5: Save all Q&A pairs for reference
    try:
        with open("all_rfp_qa.json", "w") as f:
            json.dump(all_qa_pairs, f, indent=2)
        logger.info(f"Saved {len(all_qa_pairs)} Q&A pairs to all_rfp_qa.json")
    except Exception as e:
        logger.error(f"Error saving all Q&A pairs: {e}")

    # Step 6: Prioritize and limit Q&A pairs for the final response
    prioritized_pairs = []
    priority_types = ["scope", "resource", "location", "timeline"]
    
    # First, add pairs from priority categories
    for prompt_type in priority_types:
        for pair in all_qa_pairs:
            if pair["prompt_type"] == prompt_type and len(prioritized_pairs) < 30:
                prioritized_pairs.append({
                    "question": pair["question"],
                    "answer": pair["answer"]
                })
    
    # If we still have room, add general pairs
    for pair in all_qa_pairs:
        if pair["prompt_type"] == "general" and len(prioritized_pairs) < 30:
            prioritized_pairs.append({
                "question": pair["question"],
                "answer": pair["answer"]
            })

    state["qa_pairs"] = prioritized_pairs
    logger.info(f"Selected {len(prioritized_pairs)} prioritized Q&A pairs for the response")
    return state

# Node 3: Store Q&A (for user review/UI display)
def store_qa_pairs(state: RFPState) -> RFPState:
    logger.info("Storing Q&A pairs")
    try:
        with open("rfp_qa.json", "w") as f:
            json.dump(state["qa_pairs"], f, indent=2)
        logger.info("Q&A pairs saved to rfp_qa.json")
    except Exception as e:
        logger.error(f"Error saving Q&A pairs: {e}")
    return state

# Node 4: Generate RFP response as a DOCX
def generate_response(state: RFPState) -> RFPState:
    logger.info("Generating RFP response as a DOCX")
    if not state["qa_pairs"]:
        logger.warning("No Q&A pairs available for response generation")
        state["response"] = "No response generated due to missing Q&A pairs."
        try:
            with open("rfp_response.txt", "w") as f:
                f.write(state["response"])
            logger.info("Response saved to rfp_response.txt")
        except Exception as e:
            logger.error(f"Error saving response: {e}")
        return state

    # Create a new Document
    doc = Document()

    try:
        # Title Page
        doc.add_heading('RFP Response: AI-Driven Automation Platform', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph('ApexNeural Inc.')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(16)
        p = doc.add_paragraph('Submitted to: Bidder Organization')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(14)
        p = doc.add_paragraph('May 24, 2025')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(14)
        doc.add_page_break()

        # Introduction
        doc.add_heading('Introduction', 1)
        doc.add_paragraph(
            "We, ApexNeural Inc., are pleased to submit this response to your Request for Proposal (RFP) for an AI-driven automation platform for neural network training. Our organization is committed to delivering cutting-edge AI solutions to enhance neural network training efficiency. This response outlines our approach to fulfilling the requirements set forth in your RFP, detailing our methodology, implementation plan, resource allocation, and pricing structure. We aim to establish a collaborative partnership that ensures the successful execution of this project."
        )

        # Objective Statement
        doc.add_heading('Objective Statement', 1)
        doc.add_paragraph(
            "The primary objective of this response is to demonstrate our capability to design, develop, and deploy a cloud-based automation platform that integrates with ApexNeural’s existing infrastructure (TensorFlow, PyTorch, Kubernetes) and complies with industry standards for security and performance, as specified in your RFP. Our goal is to provide a scalable, efficient, and secure solution that streamlines neural network training and supports ApexNeural’s data lakes and compute clusters."
        )

        # Readout of Requirements
        doc.add_heading('Readout of Requirements', 1)
        doc.add_paragraph(
            "Below is a summary of the key requirements extracted from your RFP, ensuring we fully understand and address your expectations:"
        )
        for pair in state["qa_pairs"]:
            if not (pair.get('question') and pair.get('answer')):
                logger.warning(f"Skipping invalid Q&A pair: {pair}")
                continue
            p = doc.add_paragraph()
            run = p.add_run(f"Q: {pair['question']}")
            run.bold = True
            p.add_run(f"\nA: {pair['answer']}")

        # Methodology
        doc.add_heading('Methodology', 1)
        doc.add_paragraph(
            "Our approach to fulfilling the RFP requirements involves a structured methodology to ensure the successful design, development, and deployment of the AI-driven automation platform. The methodology is divided into the following phases:"
        )
        methodology = [
            ("Requirement Analysis", "Collaborate with ApexNeural to understand infrastructure and requirements, ensuring compatibility with TensorFlow, PyTorch, and Kubernetes."),
            ("Platform Design", "Design a cloud-based automation platform with modular architecture to support scalability and integration with data lakes and compute clusters."),
            ("Development", "Develop the platform using agile methodologies, incorporating security features and performance optimizations as per industry standards."),
            ("Testing and Validation", "Conduct rigorous testing, including integration tests with ApexNeural’s systems and performance benchmarking."),
            ("Deployment and Maintenance", "Deploy the platform to ApexNeural’s cloud environment and provide ongoing maintenance and support.")
        ]
        for phase, description in methodology:
            p = doc.add_paragraph(f"{phase}: ", style='List Number')
            p.add_run(description)

        # Implementation Plan
        doc.add_heading('Implementation Plan', 1)
        doc.add_paragraph(
            "The implementation plan outlines the key milestones and timelines for the project, ensuring delivery within the specified timeline from the RFP."
        )
        table = doc.add_table(rows=6, cols=3)
        table.style = 'Table Grid'
        headers = ["Phase", "Activities", "Timeline"]
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
            table.cell(0, i).paragraphs[0].runs[0].font.bold = True
        timeline_data = [
            ("Requirement Analysis", "Stakeholder meetings, requirement gathering", "Months 1--2"),
            ("Platform Design", "Architecture design, integration planning", "Months 3--4"),
            ("Development", "Coding, security implementation", "Months 5--8"),
            ("Testing and Validation", "Integration and performance testing", "Months 9--10"),
            ("Deployment and Maintenance", "Platform deployment, ongoing support", "Month 11 onwards")
        ]
        for row_idx, (phase, activities, timeline) in enumerate(timeline_data, 1):
            table.cell(row_idx, 0).text = phase
            table.cell(row_idx, 1).text = activities
            table.cell(row_idx, 2).text = timeline
        doc.add_paragraph("Table: Implementation Timeline")

        # Resources Needed by Role and Phase
        doc.add_heading('Resources Needed by Role and Phase', 1)
        doc.add_paragraph(
            "The project requires a diverse team across different phases, with specific roles and responsibilities:"
        )
        resources = [
            ("Requirement Analysis (Months 1--2)", [
                "Project Manager (1): Oversees planning and coordination.",
                "Business Analyst (2): Gathers and documents requirements."
            ]),
            ("Platform Design (Months 3--4)", [
                "Solutions Architect (2): Designs platform architecture.",
                "Security Specialist (1): Ensures compliance with security standards."
            ]),
            ("Development (Months 5--8)", [
                "Software Engineers (5): Develop platform components.",
                "DevOps Engineer (2): Manages CI/CD and Kubernetes integration."
            ]),
            ("Testing and Validation (Months 9--10)", [
                "QA Engineers (3): Conducts testing and validation.",
                "Integration Specialist (1): Ensures system compatibility."
            ]),
            ("Deployment and Maintenance (Month 11 onwards)", [
                "Operations Manager (1): Oversees platform operations.",
                "Support Engineers (2): Provides ongoing maintenance."
            ])
        ]
        for phase, roles in resources:
            doc.add_paragraph(phase + ":", style='List Bullet')
            for role in roles:
                doc.add_paragraph(role, style='List Bullet 2')

        # Detailed Pricing
        doc.add_heading('Detailed Pricing', 1)
        doc.add_paragraph(
            "The pricing is broken down by resource and technology, reflecting the costs associated with each phase of the project. All figures are in INR (Indian Rupees)."
        )
        table = doc.add_table(rows=14, cols=3)
        table.style = 'Table Grid'
        headers = ["Category", "Item", "Cost (INR)"]
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
            table.cell(0, i).paragraphs[0].runs[0].font.bold = True
        pricing_data = [
            ("Human Resources", "", ""),
            ("", "Project Manager", "1,200,000"),
            ("", "Business Analyst", "800,000"),
            ("", "Solutions Architect", "1,000,000"),
            ("", "Security Specialist", "600,000"),
            ("", "Software Engineers", "3,000,000"),
            ("", "DevOps Engineer", "1,200,000"),
            ("", "QA Engineers", "900,000"),
            ("", "Integration Specialist", "400,000"),
            ("", "Operations Manager", "800,000"),
            ("", "Support Engineers", "1,000,000"),
            ("Technology and Equipment", "", ""),
            ("", "Cloud Infrastructure", "5,000,000"),
            ("", "Development Tools", "1,000,000"),
            ("", "Total Estimated Cost", "16,900,000")
        ]
        for row_idx, (category, item, cost) in enumerate(pricing_data, 1):
            table.cell(row_idx, 0).text = category
            table.cell(row_idx, 1).text = item
            table.cell(row_idx, 2).text = cost
        doc.add_paragraph("Table: Detailed Pricing Breakdown")

        # Conclusion
        doc.add_heading('Conclusion', 1)
        doc.add_paragraph(
            "We are confident that our proposed approach to the AI-driven automation platform meets the requirements outlined in your RFP. By leveraging our expertise in AI and cloud technologies, a skilled workforce, and a well-defined implementation plan, we aim to deliver a scalable and secure platform that enhances ApexNeural’s neural network training capabilities. Our competitive pricing and commitment to quality make us a strong partner for this initiative. We look forward to the opportunity to collaborate with your organization."
        )

        # Save the DOCX file
        doc.save("rfp_response.docx")
        logger.info("DOCX file 'rfp_response.docx' generated successfully")
        state["response"] = "RFP response generated as 'rfp_response.docx'."
    except Exception as e:
        logger.error(f"Error generating DOCX file: {e}")
        state["response"] = f"Error generating DOCX file: {e}"
        return state

    return state

# Define the LangGraph workflow
workflow = StateGraph(RFPState)

# Add nodes
workflow.add_node("extract_rfp", extract_rfp_text)
workflow.add_node("generate_qa", generate_qa_pairs)
workflow.add_node("store_qa", store_qa_pairs)
workflow.add_node("generate_response", generate_response)

# Define edges
workflow.add_edge("extract_rfp", "generate_qa")
workflow.add_edge("generate_qa", "store_qa")
workflow.add_edge("store_qa", "generate_response")
workflow.add_edge("generate_response", END)

# Set entry point
workflow.set_entry_point("extract_rfp")

# Compile the graph
app = workflow.compile()

# Example usage
if __name__ == "__main__":
    # Construct absolute path to apexneural_rfp.pdf
    script_dir = os.path.dirname(os.path.abspath(__file__))
    rfp_path = os.path.join(script_dir, "apexneural_rfp.pdf")
    logger.info(f"Looking for RFP file at: {rfp_path}")
    
    # Allow custom path via command-line argument
    if len(sys.argv) > 1:
        rfp_path = sys.argv[1]
        logger.info(f"Using custom RFP path from command line: {rfp_path}")
    
    initial_state = RFPState(rfp_path=rfp_path)
    
    try:
        final_state = app.invoke(initial_state)
        print("Generated Q&A Pairs:")
        for pair in final_state["qa_pairs"]:
            print(f"Q: {pair['question']}\nA: {pair['answer']}\n")
        print("Generated RFP Response:")
        print(final_state["response"])
    except Exception as e:
        logger.error(f"Workflow failed: {e}")
        print(f"Error running workflow: {e}")