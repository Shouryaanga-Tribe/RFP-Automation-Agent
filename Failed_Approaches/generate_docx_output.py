from typing import Dict, List, TypedDict
from langgraph.graph import StateGraph, END
from langchain.prompts import PromptTemplate
from langchain_community.llms import HuggingFacePipeline
from transformers import pipeline
import pdfplumber
import json
import logging
import os
import re
import sys
import pytesseract
from PIL import Image
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
import traceback

# Set up logging for debugging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

# Define the state to track data across the workflow
class RFPState(TypedDict):
    rfp_path: str
    rfp_text: str
    qa_pairs: List[Dict[str, str]]
    response: str

# Verify dependencies
try:
    import spacy
    nlp = spacy.load("en_core_web_sm")
    logger.info("spaCy model loaded successfully")
    use_spacy = True
except Exception as e:
    logger.warning(f"spaCy not available, using keyword-based extraction: {e}")
    use_spacy = False

try:
    llm = HuggingFacePipeline.from_model_id(
        model_id="google/flan-t5-small",
        task="text2text-generation",
        pipeline_kwargs={"max_length": 1024, "truncation": True}
    )
    logger.info("LLM model loaded successfully")
except Exception as e:
    logger.error(f"Failed to load LLM: {e}")
    raise

# Node 1: Extract text from RFP PDF
def extract_rfp_text(state: RFPState) -> RFPState:
    logger.info(f"Attempting to extract text from RFP: {state['rfp_path']}")
    if not os.path.exists(state["rfp_path"]):
        error_msg = (
            f"RFP file not found at: {state['rfp_path']}\n"
            "Please ensure 'apexneural_rfp.pdf' exists in the project directory."
        )
        logger.error(error_msg)
        raise FileNotFoundError(error_msg)
    
    try:
        text = ""
        with pdfplumber.open(state["rfp_path"]) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                else:
                    logger.info("No text extracted from page, attempting OCR...")
                    try:
                        page_image = page.to_image(resolution=300)
                        img_byte_arr = io.BytesIO()
                        page_image.original.save(img_byte_arr, format='PNG')
                        img_byte_arr.seek(0)
                        img = Image.open(img_byte_arr)
                        ocr_text = pytesseract.image_to_string(img)
                        if ocr_text:
                            text += ocr_text + "\n"
                            logger.info(f"OCR extracted text: {ocr_text[:100]}...")
                    except Exception as ocr_error:
                        logger.warning(f"OCR failed for page: {ocr_error}")
                        continue
        
        if not text.strip():
            logger.warning("No text extracted from RFP after attempting OCR.")
            state["rfp_text"] = ""
        else:
            logger.info(f"Extracted text: {text[:200]}...")
            state["rfp_text"] = text
            logger.info("RFP text extracted successfully")
    except Exception as e:
        logger.error(f"Error extracting RFP text: {e}")
        state["rfp_text"] = ""
    return state

# Node 2: Generate Q&A from RFP text
def generate_qa_pairs(state: RFPState) -> RFPState:
    logger.info("Generating Q&A pairs from RFP text")
    if not state["rfp_text"]:
        logger.warning("No RFP text available for Q&A generation")
        state["qa_pairs"] = []
        return state

    all_qa_pairs = []
    
    # Step 1: Extract potential requirements
    if use_spacy:
        doc = nlp(state["rfp_text"])
        requirements = [
            sent.text.strip() for sent in doc.sents
            if any(keyword in sent.text.lower() for keyword in ["must", "shall", "required", "scope", "resource", "location", "capacity", "timeline", "duration", "metrics", "evaluation"])
        ]
    else:
        sentences = re.split(r'[.!?]\s+', state["rfp_text"])
        requirements = [
            sent.strip() for sent in sentences
            if any(keyword in sent.lower() for keyword in ["must", "shall", "required", "scope", "resource", "location", "capacity", "timeline", "duration", "metrics", "evaluation"])
        ]
    
    # Step 2: Look for numbered clauses (e.g., "1.1.1") and questions (e.g., "Q6...")
    clause_pattern = r'\d+\.\d+(\.\d+)?\s+.*?(?=\n\d+\.\d+|\n[A-Z]|\Z)'
    question_pattern = r'Q\d+\s+.*?(?=\nQ\d+|\n[A-Z]|\Z)'
    clauses = re.findall(clause_pattern, state["rfp_text"], re.DOTALL)
    questions = re.findall(question_pattern, state["rfp_text"], re.DOTALL)
    requirements.extend([clause.strip() for clause in clauses if clause.strip()])
    requirements.extend([q.strip() for q in questions if q.strip()])
    
    # Remove duplicates while preserving order
    seen = set()
    requirements = [req for req in requirements if not (req in seen or seen.add(req))]
    
    logger.info(f"Found {len(requirements)} potential requirements: {requirements[:3]}...")

    # Step 3: Define enhanced prompt templates
    qa_prompts = {
        "scope": PromptTemplate(
            input_variables=["requirement"],
            template="Based on the requirement: '{requirement}'\nGenerate a question and answer about the scope of work.\nQuestion: What is the scope of work for this requirement?\nAnswer: The scope includes {requirement}"
        ),
        "resource": PromptTemplate(
            input_variables=["requirement"],
            template="Based on the requirement: '{requirement}'\nGenerate a question and answer about the resources needed.\nQuestion: What resources are required to meet this requirement?\nAnswer: The resources needed include {requirement}"
        ),
        "location": PromptTemplate(
            input_variables=["requirement"],
            template="Based on the requirement: '{requirement}'\nGenerate a question and answer about the work location.\nQuestion: Where must this work be performed?\nAnswer: The work location is specified as {requirement}"
        ),
        "timeline": PromptTemplate(
            input_variables=["requirement"],
            template="Based on the requirement: '{requirement}'\nGenerate a question and answer about the timeline.\nQuestion: What is the timeline for meeting this requirement?\nAnswer: The timeline is {requirement}"
        ),
        "metrics": PromptTemplate(
            input_variables=["requirement"],
            template="Based on the requirement: '{requirement}'\nGenerate a question and answer about performance metrics.\nQuestion: How will performance be measured for this requirement?\nAnswer: Performance will be measured by {requirement}"
        ),
        "evaluation": PromptTemplate(
            input_variables=["requirement"],
            template="Based on the requirement: '{requirement}'\nGenerate a question and answer about evaluation criteria.\nQuestion: What are the evaluation criteria for this requirement?\nAnswer: The evaluation criteria include {requirement}"
        ),
        "general": PromptTemplate(
            input_variables=["requirement"],
            template="Based on the requirement: '{requirement}'\nGenerate a question and answer.\nQuestion: What does this requirement entail?\nAnswer: This requirement entails {requirement}"
        )
    }
    
    # Step 4: Generate Q&A pairs for all requirements
    for idx, req in enumerate(requirements):
        try:
            # Clean up the requirement text
            cleaned_req = req.replace('\n', ' ').replace('-', '').strip()
            if not cleaned_req:
                continue
                
            # Check if the requirement is a question (e.g., starts with "Q")
            is_question = cleaned_req.startswith('Q')
            
            # Determine the type of prompt to use based on keywords
            prompt_type = "general"
            if any(keyword in cleaned_req.lower() for keyword in ["scope", "work", "responsibility"]):
                prompt_type = "scope"
            elif any(keyword in cleaned_req.lower() for keyword in ["resource", "cost", "security", "deposit"]):
                prompt_type = "resource"
            elif any(keyword in cleaned_req.lower() for keyword in ["location", "place", "address", "grid"]):
                prompt_type = "location"
            elif any(keyword in cleaned_req.lower() for keyword in ["timeline", "duration", "date", "period"]):
                prompt_type = "timeline"
            elif any(keyword in cleaned_req.lower() for keyword in ["metrics", "performance", "utilization"]):
                prompt_type = "metrics"
            elif any(keyword in cleaned_req.lower() for keyword in ["evaluation", "criteria"]):
                prompt_type = "evaluation"

            if is_question:
                question = cleaned_req
                prompt = PromptTemplate(
                    input_variables=["requirement"],
                    template=f"Based on the question: '{cleaned_req}'\nGenerate an answer.\nAnswer: {{requirement}}"
                ).format(requirement=cleaned_req)
                qa_text = llm.invoke(prompt)
                answer = qa_text.replace("Answer: ", "").strip() if qa_text.startswith("Answer: ") else qa_text
            else:
                prompt = qa_prompts[prompt_type].format(requirement=cleaned_req)
                qa_text = llm.invoke(prompt)
                logger.info(f"LLM output for requirement {idx + 1} ('{cleaned_req[:50]}...'): {qa_text}")
                
                # Parse the LLM output
                try:
                    if "\n" in qa_text and "Question: " in qa_text and "Answer: " in qa_text:
                        parts = qa_text.split("\n")
                        question = next((line.replace("Question: ", "").strip() for line in parts if line.startswith("Question: ")), cleaned_req)
                        answer = next((line.replace("Answer: ", "").strip() for line in parts if line.startswith("Answer: ")), cleaned_req)
                    else:
                        question = qa_prompts[prompt_type].template.split("Question: ")[1].split("\n")[0].format(requirement=cleaned_req)
                        answer = cleaned_req
                except Exception as parse_error:
                    logger.warning(f"Failed to parse Q&A for requirement {idx + 1}: {cleaned_req[:50]}... {parse_error}")
                    question = qa_prompts[prompt_type].template.split("Question: ")[1].split("\n")[0].format(requirement=cleaned_req)
                    answer = cleaned_req

            if question and answer:
                all_qa_pairs.append({
                    "question": question,
                    "answer": answer,
                    "prompt_type": prompt_type,
                    "requirement": cleaned_req
                })
                logger.info(f"Generated Q&A {idx + 1}: Q: {question[:50]}... A: {answer[:50]}...")
            else:
                logger.warning(f"Empty question or answer for requirement {idx + 1}: {cleaned_req[:50]}...")
                question = f"What does this requirement entail?"
                answer = cleaned_req
                all_qa_pairs.append({
                    "question": question,
                    "answer": answer,
                    "prompt_type": prompt_type,
                    "requirement": cleaned_req
                })
                logger.info(f"Fallback Q&A {idx + 1}: Q: {question[:50]}... A: {answer[:50]}...")
        except Exception as e:
            logger.error(f"Error generating Q&A for requirement {idx + 1}: {cleaned_req[:50]}... {e}")
            continue

    # Step 5: Save all Q&A pairs for reference
    try:
        with open("all_rfp_qa.json", "w") as f:
            json.dump(all_qa_pairs, f, indent=2)
        logger.info(f"Saved {len(all_qa_pairs)} Q&A pairs to all_rfp_qa.json")
    except Exception as e:
        logger.error(f"Error saving all Q&A pairs: {e}")

    # Step 6: Prioritize and limit Q&A pairs for the final response
    prioritized_pairs = []
    priority_types = ["scope", "resource", "location", "timeline", "metrics", "evaluation"]
    
    for prompt_type in priority_types:
        for pair in all_qa_pairs:
            if pair["prompt_type"] == prompt_type and len(prioritized_pairs) < 30:
                prioritized_pairs.append({
                    "question": pair["question"],
                    "answer": pair["answer"]
                })
    
    for pair in all_qa_pairs:
        if pair["prompt_type"] == "general" and len(prioritized_pairs) < 30:
            prioritized_pairs.append({
                "question": pair["question"],
                "answer": pair["answer"]
            })

    state["qa_pairs"] = prioritized_pairs
    logger.info(f"Selected {len(prioritized_pairs)} prioritized Q&A pairs for the response")
    return state

# Node 3: Store Q&A (for user review/UI display)
def store_qa_pairs(state: RFPState) -> RFPState:
    logger.info("Storing Q&A pairs")
    try:
        with open("rfp_qa.json", "w") as f:
            json.dump(state["qa_pairs"], f, indent=2)
        logger.info("Q&A pairs saved to rfp_qa.json")
    except Exception as e:
        logger.error(f"Error saving Q&A pairs: {e}")
    return state

# Node 4: Generate RFP response as a DOCX
def generate_response(state: RFPState) -> RFPState:
    logger.info("Generating RFP response as a DOCX")
    if not state["qa_pairs"]:
        logger.warning("No Q&A pairs available for response generation")
        state["response"] = "No response generated due to missing Q&A pairs."
        try:
            with open("rfp_response.txt", "w") as f:
                f.write(state["response"])
            logger.info("Response saved to rfp_response.txt")
        except Exception as e:
            logger.error(f"Error saving response: {e}")
        return state

    doc = Document()
    try:
        # Title Page
        doc.add_heading('RFP Response: AI-Driven Automation Platform', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph('ApexNeural Inc.')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(16)
        p = doc.add_paragraph('Submitted to: Bidder Organization')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(14)
        p = doc.add_paragraph('May 24, 2025')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(14)
        doc.add_page_break()

        # Introduction
        doc.add_heading('Introduction', 1)
        doc.add_paragraph(
            "ApexNeural Inc. is pleased to submit this response to your Request for Proposal (RFP) for an AI-driven automation platform for neural network training. As a leader in artificial intelligence, we are committed to delivering a scalable, secure, and efficient platform that integrates seamlessly with your existing infrastructure (TensorFlow, PyTorch, Kubernetes). This response outlines our approach, including methodology, implementation plan, resource allocation, and pricing, to meet your requirements and ensure project success."
        )

        # Objective Statement
        doc.add_heading('Objective Statement', 1)
        doc.add_paragraph(
            "Our objective is to develop and deploy a cloud-based automation platform that streamlines neural network training, integrates with ApexNeural’s data lakes and compute clusters, and complies with industry standards for security and performance. We aim to deliver a solution within the specified timeline, ensuring high performance and scalability to support your AI initiatives."
        )

        # Readout of Requirements
        doc.add_heading('Readout of Requirements', 1)
        doc.add_paragraph(
            "Below is a summary of the key requirements extracted from your RFP, with our responses to ensure alignment with your expectations:"
        )
        for pair in state["qa_pairs"]:
            if not (pair.get('question') and pair.get('answer')):
                logger.warning(f"Skipping invalid Q&A pair: {pair}")
                continue
            p = doc.add_paragraph()
            run = p.add_run(f"Q: {pair['question']}")
            run.bold = True
            p.add_run(f"\nA: {pair['answer']}")

        # Methodology
        doc.add_heading('Methodology', 1)
        doc.add_paragraph(
            "Our methodology ensures the successful design, development, and deployment of the AI-driven automation platform, with the following phases:"
        )
        methodology = [
            ("Requirement Analysis", "Engage with ApexNeural to validate requirements, focusing on integration with TensorFlow, PyTorch, and Kubernetes, and compliance with security standards."),
            ("System Design", "Design a modular, cloud-based platform architecture to ensure scalability, security, and compatibility with ApexNeural’s data lakes and compute clusters."),
            ("Development", "Implement the platform using agile development practices, incorporating robust security measures and performance optimizations."),
            ("Testing and Validation", "Perform comprehensive testing, including integration, security, and performance tests, to ensure reliability and compliance."),
            ("Deployment and Support", "Deploy the platform to ApexNeural’s cloud environment and provide ongoing maintenance and support.")
        ]
        for phase, description in methodology:
            p = doc.add_paragraph(f"{phase}: ", style='List Number')
            p.add_run(description)

        # Implementation Plan
        doc.add_heading('Implementation Plan', 1)
        doc.add_paragraph(
            "The implementation plan outlines key milestones and timelines to deliver the platform within the specified period, as per the RFP."
        )
        timeline_data = [
            ("Requirement Analysis", "Stakeholder meetings, requirement validation", "Months 1-2"),
            ("System Design", "Architecture design, integration planning", "Months 3-4"),
            ("Development", "Platform development, security implementation", "Months 5-8"),
            ("Testing and Validation", "Integration and performance testing", "Months 9-10"),
            ("Deployment and Support", "Platform deployment, ongoing support", "Month 11 onwards")
        ]
        table = doc.add_table(rows=len(timeline_data) + 1, cols=3)
        table.style = 'Table Grid'
        headers = ["Phase", "Activities", "Timeline"]
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
        for row_idx, (phase, activities, timeline) in enumerate(timeline_data, 1):
            table.cell(row_idx, 0).text = phase
            table.cell(row_idx, 1).text = activities
            table.cell(row_idx, 2).text = timeline
        doc.add_paragraph("Table: Implementation Timeline")

        # Resources Needed by Role and Phase
        doc.add_heading('Resources Needed by Role and Phase', 1)
        doc.add_paragraph(
            "The project requires a skilled team across various phases, with specific roles to ensure success:"
        )
        resources = [
            ("Requirement Analysis (Months 1-2)", [
                "Project Manager (1): Oversees project planning and coordination.",
                "Business Analyst (2): Validates and documents requirements."
            ]),
            ("System Design (Months 3-4)", [
                "Solutions Architect (2): Designs platform architecture.",
                "Security Specialist (1): Ensures compliance with security standards."
            ]),
            ("Development (Months 5-8)", [
                "Software Engineers (5): Develops platform components.",
                "DevOps Engineer (2): Manages CI/CD pipelines and Kubernetes integration."
            ]),
            ("Testing and Validation (Months 9-10)", [
                "QA Engineers (3): Conducts integration and performance testing.",
                "Integration Specialist (1): Ensures compatibility with existing systems."
            ]),
            ("Deployment and Support (Month 11 onwards)", [
                "Operations Manager (1): Oversees platform operations.",
                "Support Engineers (2): Provides maintenance and technical support."
            ])
        ]
        for phase, roles in resources:
            doc.add_paragraph(phase + ":", style='List Bullet')
            for role in roles:
                doc.add_paragraph(role, style='List Bullet 2')

        # Detailed Pricing
        doc.add_heading('Detailed Pricing', 1)
        doc.add_paragraph(
            "The pricing reflects the costs for human resources and technology, in INR (Indian Rupees)."
        )
        pricing_data = [
            ("Human Resources", "", ""),
            ("", "Project Manager (1, 12 months)", "1,200,000"),
            ("", "Business Analyst (2, 2 months)", "800,000"),
            ("", "Solutions Architect (2, 2 months)", "1,000,000"),
            ("", "Security Specialist (1, 2 months)", "600,000"),
            ("", "Software Engineers (5, 4 months)", "3,000,000"),
            ("", "DevOps Engineer (2, 4 months)", "1,200,000"),
            ("", "QA Engineers (3, 2 months)", "900,000"),
            ("", "Integration Specialist (1, 2 months)", "400,000"),
            ("", "Operations Manager (1, 12 months)", "800,000"),
            ("", "Support Engineers (2, 12 months)", "1,000,000"),
            ("Technology and Equipment", "", ""),
            ("", "Cloud Infrastructure (AWS/GCP)", "5,000,000"),
            ("", "Development Tools and Licenses", "1,000,000"),
            ("", "Total Estimated Cost", "16,900,000")
        ]
        table = doc.add_table(rows=len(pricing_data) + 1, cols=3)
        table.style = 'Table Grid'
        headers = ["Category", "Item", "Cost (INR)"]
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
        for row_idx, (category, item, cost) in enumerate(pricing_data, 1):
            table.cell(row_idx, 0).text = category
            table.cell(row_idx, 1).text = item
            table.cell(row_idx, 2).text = cost
        doc.add_paragraph("Table: Detailed Pricing Breakdown")

        # Conclusion
        doc.add_heading('Conclusion', 1)
        doc.add_paragraph(
            "ApexNeural Inc. is confident in delivering an AI-driven automation platform that meets your RFP requirements. Our solution ensures seamless integration with TensorFlow, PyTorch, and Kubernetes, robust security, and high performance. With a skilled team and competitive pricing, we are committed to a successful partnership. We look forward to collaborating with you to advance your AI initiatives."
        )

        # Save the DOCX file
        doc.save("rfp_response.docx")
        logger.info("DOCX file 'rfp_response.docx' generated successfully")
        state["response"] = "RFP response generated as 'rfp_response.docx'."
    except Exception as e:
        logger.error(f"Error generating DOCX file: {e}\n{traceback.format_exc()}")
        state["response"] = f"Error generating DOCX file: {e}"
        return state

    return state

# Define the LangGraph workflow
workflow = StateGraph(RFPState)

# Add nodes
workflow.add_node("extract_rfp", extract_rfp_text)
workflow.add_node("generate_qa", generate_qa_pairs)
workflow.add_node("store_qa", store_qa_pairs)
workflow.add_node("generate_response", generate_response)

# Define edges
workflow.add_edge("extract_rfp", "generate_qa")
workflow.add_edge("generate_qa", "store_qa")
workflow.add_edge("store_qa", "generate_response")
workflow.add_edge("generate_response", END)

# Set entry point
workflow.set_entry_point("extract_rfp")

# Compile the graph
app = workflow.compile()

# Example usage
if __name__ == "__main__":
    # Construct absolute path to apexneural_rfp.pdf
    script_dir = os.path.dirname(os.path.abspath(__file__))
    rfp_path = os.path.join(script_dir, "18pages.pdf")
    logger.info(f"Looking for RFP file at: {rfp_path}")
    
    # Allow custom path via command-line argument
    if len(sys.argv) > 1:
        rfp_path = sys.argv[1]
        logger.info(f"Using custom RFP path from command line: {rfp_path}")
    
    initial_state = RFPState(rfp_path=rfp_path)
    
    try:
        final_state = app.invoke(initial_state)
        print("Generated Q&A Pairs:")
        for pair in final_state["qa_pairs"]:
            print(f"Q: {pair['question']}\nA: {pair['answer']}\n")
        print("Generated RFP Response:")
        print(final_state["response"])
    except Exception as e:
        logger.error(f"Workflow failed: {e}\n{traceback.format_exc()}")
        print(f"Error running workflow: {e}")