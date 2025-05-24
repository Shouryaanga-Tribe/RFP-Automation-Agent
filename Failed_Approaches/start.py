import pypdf
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import spacy
import re
import os
from typing import List, Tuple

# Load spaCy model
nlp = spacy.load("en_core_web_sm")
    
def read_pdf(file_path: str) -> str:
    """Read text from a PDF file using pypdf."""
    try:
        with open(file_path, 'rb') as file:
            reader = pypdf.PdfReader(file)
            text = ""
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
            if not text.strip():
                return "Error: No text could be extracted from the PDF."
            return text
    except Exception as e:
        return f"Error reading PDF: {str(e)}"

def read_docx(file_path: str) -> str:
    """Read text from a DOCX file."""
    try:
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        if not text.strip():
            return "Error: No text could be extracted from the DOCX."
        return text
    except Exception as e:
        return f"Error reading DOCX: {str(e)}"

def extract_questions(text: str) -> List[Tuple[str, str]]:
    """Extract questions from RFP text using spaCy for better NLP."""
    try:
        doc = nlp(text)
        questions = []
        requirement_keywords = ['require', 'must', 'should', 'provide', 'submit', 'include', 'describe', 'explain']

        for sent in doc.sents:
            sent_text = sent.text.strip()
            # Direct questions
            if sent_text.endswith('?'):
                questions.append((sent_text, "Please provide your response here."))
            # Requirements converted to questions
            elif any(keyword in sent_text.lower() for keyword in requirement_keywords):
                question = f"What {sent_text.lower().replace('the proposer', 'is your').replace('must', '').replace('should', '').strip()}?"
                question = re.sub(r'\s+', ' ', question).strip()
                if not question.endswith('?'):
                    question += '?'
                questions.append((question, "Please provide your response here."))
        
        return questions if questions else [("No questions or requirements identified.", "N/A")]
    except Exception as e:
        return [(f"Error extracting questions: {str(e)}", "N/A")]

def generate_rfp_response(questions_answers: List[Tuple[str, str]]) -> Document:
    """Generate a professional RFP response document with enhanced formatting."""
    doc = Document()

    # Add company logo placeholder
    doc.add_paragraph("[Insert Company Logo Here]", style='Title').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('Request for Proposal (RFP) Response', 0)
    doc.add_paragraph('Prepared by: [Your Company Name]')
    doc.add_paragraph('Address: [Your Company Address]')
    doc.add_paragraph('Date: May 24, 2025')
    doc.add_paragraph()

    # Add table of contents
    doc.add_heading('Table of Contents', 1)
    doc.add_paragraph('1. Introduction', style='List Number')
    for i, _ in enumerate(questions_answers, 2):
        doc.add_paragraph(f"{i}. Response to Question {i-1}", style='List Number')
    doc.add_page_break()

    # Introduction
    doc.add_heading('1. Introduction', 1)
    doc.add_paragraph(
        "We are pleased to submit our response to your Request for Proposal (RFP). "
        "Our team has carefully reviewed the requirements and is confident in our ability "
        "to deliver a solution that meets your needs."
    )
    doc.add_paragraph()

    # Responses
    doc.add_heading('Responses to RFP Requirements', 1)
    for i, (question, answer) in enumerate(questions_answers, 1):
        doc.add_heading(f"{i+1}. Question {i}: {question}", level=2)
        doc.add_paragraph(answer)
        doc.add_paragraph()

    return doc

def save_response(doc: Document, output_path: str) -> str:
    """Save the RFP response to a DOCX file."""
    try:
        doc.save(output_path)
        return f"Response saved to {output_path}"
    except Exception as e:
        return f"Error saving response: {str(e)}"

def main():
    """Main function to run the RFP automation agent."""
    file_path = input("Enter the path to the RFP document (PDF or DOCX): ").strip()
    
    # Validate file existence and format
    if not os.path.exists(file_path):
        print(f"Error: File '{file_path}' does not exist.")
        return
    if not (file_path.lower().endswith('.pdf') or file_path.lower().endswith('.docx')):
        print("Unsupported file format. Please use PDF or DOCX.")
        return

    # Phase 1: Read and extract questions
    if file_path.lower().endswith('.pdf'):
        text = read_pdf(file_path)
    else:
        text = read_docx(file_path)

    if "Error" in text:
        print(text)
        return

    questions = extract_questions(text)
    if not questions or questions[0][0].startswith("Error"):
        print(questions[0][0])
        return

    # Display questions and collect answers
    print("\nExtracted Questions:")
    user_answers = []
    for i, (question, _) in enumerate(questions, 1):
        print(f"Question {i}: {question}")
        answer = input(f"Provide your answer for Question {i} (or press Enter to skip): ").strip()
        user_answers.append((question, answer or "Please provide your response here."))

    # Phase 2: Generate and save response
    response_doc = generate_rfp_response(user_answers)
    output_path = "rfp_response.docx"
    result = save_response(response_doc, output_path)
    print(result)

if __name__ == "__main__":
    main()